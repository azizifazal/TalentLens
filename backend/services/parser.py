from __future__ import annotations

import io
from typing import Any

import boto3
import structlog
from pypdf import PdfReader

from backend.core.config import get_settings
from backend.core.exceptions import ParseError
from backend.models.candidate import (
    CandidateProfile,
    Certification,
    Education,
    RoleLevel,
    Skill,
    SkillDepth,
    WorkHistory,
)
from backend.services.bedrock import BedrockClient

logger = structlog.get_logger(__name__)

_PARSE_SYSTEM = """You are an expert resume parser and career analyst.
Extract structured candidate data AND behavioral evidence sentences.
Return ONLY valid JSON matching the exact schema provided.
Never invent contact information such as email or phone.
For raw_behavioral_evidence, copy exact sentences or short phrases from the resume
that reveal behavioral patterns: promotions, skill acquisition, leadership moments,
scope expansion, self-directed projects, cross-functional ownership.
If a field is unavailable, use an empty string or empty list."""

_PARSE_PROMPT_TEMPLATE = """Parse this resume and return valid JSON only.

Resume Text:
{resume_text}

Return this exact JSON schema (no extra keys, no markdown):
{{
  "full_name": "string",
  "current_title": "string",
  "current_company": "string",
  "location": "string",
  "years_experience": number,
  "skills": [
    {{
      "name": "string",
      "category": "string (e.g. Backend, Frontend, ML, DevOps, Data, Management, Soft Skills)",
      "last_used_year": number_or_null,
      "depth": "AWARE|PRACTICED|EXPERT"
    }}
  ],
  "work_history": [
    {{
      "title": "string",
      "company": "string",
      "start_date": "YYYY-MM",
      "end_date": "YYYY-MM or null if current",
      "duration_months": number,
      "level_inferred": "JUNIOR|MID|SENIOR|LEAD|PRINCIPAL|MANAGER|DIRECTOR",
      "description_summary": "string (max 200 chars)",
      "responsibilities": ["string", "string"]
    }}
  ],
  "education": [
    {{
      "degree": "string",
      "field": "string",
      "institution": "string",
      "graduation_year": number_or_null
    }}
  ],
  "certifications": [
    {{
      "name": "string",
      "issuer": "string",
      "year": number
    }}
  ],
  "raw_behavioral_evidence": [
    "exact sentence or phrase showing behavioral pattern (3-8 items)"
  ]
}}"""


class ParserService:
    def __init__(self) -> None:
        settings = get_settings()
        self._s3 = boto3.client("s3", region_name=settings.aws_region)
        self._bucket = settings.s3_bucket_name
        self._bedrock = BedrockClient()

    def download_and_extract_text(self, s3_key: str) -> str:
        try:
            response = self._s3.get_object(Bucket=self._bucket, Key=s3_key)
            file_bytes = response["Body"].read()
        except Exception as exc:
            raise ParseError(f"Failed to download file from S3: {exc}")

        file_name = s3_key.split("/")[-1].lower()
        if file_name.endswith(".pdf"):
            return self._extract_pdf_text(file_bytes)
        if file_name.endswith(".docx"):
            return self._extract_docx_text(file_bytes)
        raise ParseError(f"Unsupported file type for key: {s3_key}")

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages: list[str] = []
            for page in reader.pages:
                text = page.extract_text() or ""
                pages.append(text)
            full_text = "\n".join(pages).strip()
            if len(full_text) < 50:
                raise ParseError("PDF contains insufficient extractable text (possibly scanned)")
            return full_text
        except ParseError:
            raise
        except Exception as exc:
            raise ParseError(f"PDF extraction failed: {exc}")

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        try:
            import docx  # python-docx

            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            full_text = "\n".join(paragraphs).strip()
            if len(full_text) < 50:
                raise ParseError("DOCX contains insufficient text content")
            return full_text
        except ParseError:
            raise
        except Exception as exc:
            raise ParseError(f"DOCX extraction failed: {exc}")

    def parse_resume_with_ai(self, resume_text: str) -> CandidateProfile:
        truncated = resume_text[:12000]
        prompt = _PARSE_PROMPT_TEMPLATE.format(resume_text=truncated)
        try:
            raw = self._bedrock.invoke_claude_json(
                user_prompt=prompt,
                system_prompt=_PARSE_SYSTEM,
                max_tokens=3000,
            )
        except Exception as exc:
            raise ParseError(f"AI parsing failed: {exc}")

        return self._build_profile(raw)

    def _build_profile(self, raw: dict[str, Any]) -> CandidateProfile:
        skills: list[Skill] = []
        for s in raw.get("skills", []):
            try:
                skills.append(
                    Skill(
                        name=str(s.get("name", "")),
                        category=str(s.get("category", "")),
                        last_used_year=s.get("last_used_year"),
                        depth=SkillDepth(s.get("depth", SkillDepth.PRACTICED)),
                    )
                )
            except Exception:
                continue

        work_history: list[WorkHistory] = []
        for w in raw.get("work_history", []):
            try:
                duration = int(w.get("duration_months", 0))
                if duration <= 0:
                    duration = self._estimate_duration(w.get("start_date", ""), w.get("end_date"))
                level_raw = str(w.get("level_inferred", "MID")).upper()
                try:
                    level = RoleLevel(level_raw)
                except ValueError:
                    level = RoleLevel.MID
                work_history.append(
                    WorkHistory(
                        title=str(w.get("title", "")),
                        company=str(w.get("company", "")),
                        start_date=str(w.get("start_date", "2020-01")),
                        end_date=w.get("end_date") or None,
                        duration_months=duration,
                        level_inferred=level,
                        description_summary=str(w.get("description_summary", ""))[:200],
                        responsibilities=list(w.get("responsibilities", []))[:10],
                    )
                )
            except Exception:
                continue

        education: list[Education] = []
        for e in raw.get("education", []):
            try:
                education.append(
                    Education(
                        degree=str(e.get("degree", "")),
                        field=str(e.get("field", "")),
                        institution=str(e.get("institution", "")),
                        graduation_year=e.get("graduation_year"),
                    )
                )
            except Exception:
                continue

        certifications: list[Certification] = []
        for c in raw.get("certifications", []):
            try:
                certifications.append(
                    Certification(
                        name=str(c.get("name", "")),
                        issuer=str(c.get("issuer", "")),
                        year=int(c.get("year", 2020)),
                    )
                )
            except Exception:
                continue

        behavioral_evidence = [str(e)[:300] for e in raw.get("raw_behavioral_evidence", []) if e][
            :8
        ]

        years_exp = raw.get("years_experience", 0)
        try:
            years_exp = float(years_exp)
        except (TypeError, ValueError):
            years_exp = float(sum(w.duration_months for w in work_history)) / 12

        return CandidateProfile(
            full_name=str(raw.get("full_name", "")),
            current_title=str(raw.get("current_title", "")),
            current_company=str(raw.get("current_company", "")),
            location=str(raw.get("location", "")),
            years_experience=round(years_exp, 1),
            skills=skills,
            work_history=work_history,
            education=education,
            certifications=certifications,
            raw_behavioral_evidence=behavioral_evidence,
        )

    @staticmethod
    def _estimate_duration(start_date: str, end_date: str | None) -> int:
        try:
            start_parts = start_date.split("-")
            start_year, start_month = int(start_parts[0]), int(start_parts[1])
            if end_date:
                end_parts = end_date.split("-")
                end_year, end_month = int(end_parts[0]), int(end_parts[1])
            else:
                from datetime import datetime

                now = datetime.now()
                end_year, end_month = now.year, now.month
            return max(1, (end_year - start_year) * 12 + (end_month - start_month))
        except Exception:
            return 12
